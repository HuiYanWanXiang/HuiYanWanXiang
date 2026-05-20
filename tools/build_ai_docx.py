from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "AI协同流程说明文档_交付稿.md"
OUT = ROOT / "AI协同流程说明文档_交付稿.docx"


def set_page_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.gutter = Cm(1)



def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")



def format_paragraph(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(7)
    fmt.space_after = Pt(7)



def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    if level == 1:
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(15)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(p)



def add_body_line(doc: Document, line: str) -> None:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(p)



def build_docx() -> None:
    if not SRC.exists():
        raise FileNotFoundError(f"source markdown not found: {SRC}")

    lines = SRC.read_text(encoding="utf-8").splitlines()

    doc = Document()
    set_page_style(doc)
    set_normal_style(doc)

    for raw in lines:
        line = raw.rstrip()

        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            continue

        # Remove markdown separators and quote markers for a clean formal document.
        if line.strip() == "---":
            continue
        if line.startswith("> "):
            line = line[2:]

        add_body_line(doc, line)

    doc.save(OUT)
    print(f"Generated: {OUT}")


if __name__ == "__main__":
    build_docx()
